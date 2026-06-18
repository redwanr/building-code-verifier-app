"""One-off: build a plain-language .docx explaining how the MVP works.
Run: .venv/bin/python make_explainer_doc.py
ponytail: throwaway generator, not app code; python-docx not in requirements.txt.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- base styling ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def callout(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(0x99, 0x33, 0x00)


# ===== TITLE =====
title = doc.add_heading("RAJUK Permit-Sheet Code Verifier", level=0)
sub = doc.add_paragraph()
sr = sub.add_run("How the MVP Works — A Plain-Language Guide for Non-Technical Partners")
sr.italic = True
sr.font.size = Pt(13)
doc.add_paragraph("Version: MVP · Audience: business partners, investors, "
                  "non-engineers · No technical background needed.")
doc.add_paragraph()

# ===== 1. ONE LINE =====
h("In one sentence", 1)
para("You upload a Dhaka building-approval drawing, and the tool gives a "
     "qualified architect or engineer a fast, first-pass list of possible "
     "building-code problems to review — so a human expert spends their time "
     "checking flagged items instead of hunting through the whole drawing by hand.")
callout("Important: it is a helper for an expert, not a replacement for one. "
        "It never approves or certifies a building. A licensed professional "
        "always makes the final call.")

# ===== 2. THE PROBLEM =====
h("The problem we solve", 1)
para("In Dhaka, every building needs approval from RAJUK. The approval drawing "
     "(the \"permit sheet\") is dense — floor plans, elevations, and tables full "
     "of numbers. A professional must manually check dozens of rules: Is the floor "
     "area within the allowed limit? Are there enough exit stairs? Is fire safety "
     "provided? This is slow, repetitive, and easy to get wrong when tired.")
para("Two kinds of rules apply, from two completely separate rulebooks:")
bullet("Planning rules (from RAJUK): how big the building can be, parking, "
       "setbacks (gaps from the plot edge).")
bullet("Life-safety rules (from BNBC, the national building code): exit stairs, "
       "fire protection, room sizes.")
para("Our tool reads the drawing, checks the table-readable rules automatically, "
     "and hands the expert a tidy, triaged list.")

# ===== 3. THE USER EXPERIENCE =====
h("What using it actually looks like", 1)
para("Five simple steps on a web page:")
numbered("Log in. A shared password protects the tool (MVP-level access control).")
numbered("Upload the drawing. The user drops in the approval PDF.")
numbered("Review what the tool read. The tool shows every number it pulled off "
         "the drawing, with a confidence score. Anything it is unsure about is "
         "highlighted for the human to confirm or correct.")
numbered("Fill in a few local limits. Some limits depend on the specific plot "
         "and neighbourhood (for example, the permitted floor-area ratio for that "
         "ward). The drawing can't tell us these, so the reviewer types them in.")
numbered("Get findings + export. The tool runs its checks and groups the results "
         "into three buckets, then lets the user download a report (Markdown or a "
         "printable HTML/PDF).")

# ===== 4. UNDER THE HOOD =====
h("Under the hood: the journey of one drawing", 1)
para("Here is what actually happens, step by step, in everyday language.")

h("Step 1 — The locked front door", 2)
para("The tool sits behind a password. No password, no access. This keeps "
     "confidential client drawings away from strangers.")

h("Step 2 — Turn the PDF into a picture", 2)
para("The approval drawings are \"flattened\" — they are essentially scans or "
     "images, with no hidden, selectable text inside. So the computer cannot just "
     "copy-paste the numbers. Instead, the tool converts each PDF page into a "
     "high-resolution image (like taking a sharp photo of the sheet). A free tool "
     "called Poppler does this conversion.")

h("Step 3 — The AI \"looks\" at the drawing", 2)
para("This is the heart of it. The page-image is sent to a vision AI — a large AI "
     "model (we use Claude, made by Anthropic, by default; Google's Gemini is an "
     "alternative) that can actually see and understand images, not just text.")
para("We give the AI a precise instruction sheet: \"Look at this drawing. Find "
     "these specific values — building height, number of storeys, plot area, floor "
     "area, claimed floor-area ratio, number of exit stairs, number of lifts, "
     "parking spaces, and so on. Read only what is actually printed. Never guess. "
     "For each value, tell me how confident you are (0 to 100%), and roughly where "
     "on the sheet you found it.\"")
para("The AI replies in a strict, structured form (a list of values), which the "
     "tool then processes. It also saves a small cropped picture of the exact spot "
     "each number came from — so a human can click and verify \"yes, that 4.25 "
     "really is the FAR on the table.\"")
callout("Key point: the AI reads the drawing the way a person would — by looking — "
        "not by extracting hidden text (there is none). This is why a vision AI is "
        "essential, not ordinary text software.")

h("Step 4 — The human checkpoint (the safety gate)", 2)
para("Every value comes with a confidence score. If the AI is less than 70% "
     "confident — because a number is blurry, ambiguous, or crowded — that value "
     "is flagged. The tool will NOT use a flagged value automatically. A human "
     "must look at it and confirm or fix it first.")
para("This is deliberate. We would rather stop and ask than quietly trust a "
     "shaky reading on a safety-critical drawing.")

h("Step 5 — The rulebook check", 2)
para("Now the tool compares the confirmed numbers against the building rules. "
     "Crucially, the rules are NOT written into the program's code. They live in "
     "separate, human-readable rule files (one set for RAJUK planning, one set for "
     "BNBC life-safety). Each rule is a simple statement like \"a building over 20 "
     "metres tall must have at least two exit stairs.\"")
para("A small \"rule engine\" reads these files and applies each rule to the "
     "numbers from the drawing.")

h("Step 6 — Three buckets of findings", 2)
para("Every check lands in exactly one of three buckets:")
bullet("Likely violation (red): the numbers appear to break a rule.")
bullet("Needs verification (yellow): the tool can't be sure — maybe a value is "
       "missing, or the rule's exact threshold is still being finalised — so a "
       "human must check.")
bullet("Appears compliant (green): the numbers look fine. Note the careful "
       "wording: \"appears\", never \"is\". The tool never declares final "
       "compliance.")
para("If a rule needs a number we don't have, it reports \"cannot evaluate\" — it "
     "never silently assumes everything is fine. A missed violation is the worst "
     "thing that could happen, so the tool errs toward flagging.")

h("Step 7 — The report", 2)
para("Finally, the reviewer can mark each finding as accepted or dismissed, add "
     "notes, and download a clean report. Every report carries a disclaimer that "
     "it is decision-support only, not a certification.")

# ===== 5. TWO RULEBOOKS =====
h("Why two separate rulebooks matter", 1)
para("This is a non-negotiable design rule. Planning rules (RAJUK) and "
     "life-safety rules (BNBC) come from different authorities and must never be "
     "mixed. A planning number is never checked against a safety rule, or vice "
     "versa. Every single finding clearly states which rulebook it came from. "
     "Keeping them cleanly separated is essential for the tool to be trustworthy "
     "and legally sensible.")

# ===== 6. RULES AS DATA =====
h("Why \"rules are data, not code\" is a big deal (for the business)", 1)
para("Building rules in Bangladesh are changing right now (DAP 2025 revision, "
     "draft Building Rules 2025). If our rules were buried in program code, every "
     "change would need a software developer and a new release — slow and "
     "expensive.")
para("Instead, the rules live in plain editable files. A domain expert (architect "
     "/ planner) can add or update a rule or a threshold WITHOUT any programming. "
     "Each rulebook is stamped with its source and effective date, so we always "
     "know which version of the law we checked against. This makes the product "
     "cheap to keep current and easy to defend.")

# ===== 7. SAFETY =====
h("Our safety philosophy", 1)
para("Because this touches building safety, we built in caution at every turn:")
bullet("Decision-support only — never certification. A licensed human always "
       "decides.")
bullet("\"Appears compliant\", never \"is compliant\".")
bullet("Missing or unconfirmed data produces \"cannot evaluate\", never a silent "
       "pass.")
bullet("Low-confidence readings must be human-confirmed before any check runs.")
bullet("We optimise to catch every possible safety violation, even at the cost of "
       "some false alarms. A false alarm wastes a few minutes; a missed violation "
       "could be dangerous.")
bullet("Thresholds we are not 100% sure of are marked [VERIFY] so the expert "
       "double-checks them.")

# ===== 8. SCOPE =====
h("What the MVP deliberately does NOT do", 1)
para("To ship something trustworthy fast, we kept the scope narrow. The MVP "
     "currently does not:")
bullet("Measure geometry off the drawing pixels (e.g. measuring a corridor width "
       "with a virtual ruler) — those checks say \"needs verification\" instead of "
       "guessing.")
bullet("Do structural or earthquake review.")
bullet("Handle multi-sheet drawing sets — one sheet at a time.")
bullet("Cover building types beyond standard residential.")
bullet("Provide user accounts, billing, or permissions beyond the shared "
       "password.")
para("It focuses on a handful of trusted, table-readable checks done well: "
     "egress (exits), fire-protection presence, floor-area ratio, ground "
     "coverage, parking, and lifts.")

# ===== 9. PRIVACY =====
h("Privacy and confidentiality", 1)
para("Client drawings are confidential. So:")
bullet("Uploaded files and crops are kept only for the working session — there is "
       "no permanent archive in the MVP.")
bullet("Pages are sent to the AI provider on a \"no-training\" tier, meaning the "
       "provider does not learn from or keep the drawings, and users are told this "
       "up front.")
bullet("The AI access keys live safely on our server, never in the user's browser "
       "or in our public code.")

# ===== 10. HONEST STATUS =====
h("Honest status and current limitations", 1)
para("Where things stand today:")
bullet("The MVP is built and working end-to-end on real drawings.")
bullet("Some code thresholds are still provisional (marked [VERIFY]) until a "
       "domain expert confirms the exact 2025/2026 values.")
bullet("The AI is not perfectly consistent: reading the same sheet twice can "
       "occasionally give slightly different numbers. This is exactly why the "
       "human-confirmation step exists, and it is an area we are still tightening "
       "(for example, reading each value more than once and reconciling).")
bullet("We have not yet measured accuracy against a large set of expert-labelled "
       "drawings — that benchmarking is a near-term next step.")

# ===== GLOSSARY =====
h("Mini-glossary", 1)


def gl(term, meaning):
    p = doc.add_paragraph()
    p.add_run(term + " — ").bold = True
    p.add_run(meaning)


gl("RAJUK", "the authority that approves buildings in the Dhaka region; source of "
   "planning rules.")
gl("BNBC", "Bangladesh National Building Code; source of life-safety rules.")
gl("Permit sheet", "the building-approval drawing submitted for approval.")
gl("FAR (Floor-Area Ratio)", "total floor area divided by plot area — a key limit "
   "on how much you can build.")
gl("MGC / Ground coverage", "how much of the plot the building footprint covers.")
gl("Setback", "the required gap between the building and the plot boundary.")
gl("Egress", "the means of getting out — exit stairs and doors.")
gl("Vision AI", "an AI model that can see and understand images, not just text.")
gl("Confidence score", "the AI's own honest estimate of how sure it is about a "
   "reading.")
gl("Rule engine", "the small part of the program that applies the rule files to "
   "the numbers.")

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run("This document explains the MVP at a conceptual level for a "
                  "non-technical audience. It is not a certification of the "
                  "product and does not constitute engineering or legal advice.")
fr.italic = True
fr.font.size = Pt(9)

out = "docs/How-The-MVP-Works.docx"
doc.save(out)
print("wrote", out)
